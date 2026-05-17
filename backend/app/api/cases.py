from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import Response
from pydantic import BaseModel
from datetime import datetime, timezone
from typing import Optional, List
import uuid
import hashlib
import json
from io import BytesIO
from app.services.cloudinary_service import cloudinary_service
from app.core.authz import require_roles, get_current_user
from app.core.roles import UserRole
from app.core.config import settings
from app.services.mongo_storage import mongo_storage
from app.services.encryption_service import encryption_service
from app.api.websocket import notify_case_update

# Blockchain functions removed - using simple stubs
def approve_multisig(*args, **kwargs): return None
def get_approval_status(*args, **kwargs): return {}
def file_appeal(*args, **kwargs): return None
def get_appeal(*args, **kwargs): return {}
def deposit_escrow(*args, **kwargs): return None
def release_fine(*args, **kwargs): return None
def get_escrow(*args, **kwargs): return {}
def update_case_timeline(*args, **kwargs): return None

# PDF libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("Warning: reportlab not installed. PDF download disabled. Run: pip install reportlab")

# DOCX libraries
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX download disabled. Run: pip install python-docx")

router = APIRouter(prefix="/cases", tags=["Cases"])

class CaseCreateRequest(BaseModel):
    fir_id: str
    title: str
    description: str
    priority: str = "MEDIUM"

class CaseResponse(BaseModel):
    case_id: str
    case_number: str
    fir_id: str
    title: str
    description: str
    status: str
    priority: str
    investigator_email: str
    created_at: str
    updated_at: str
    ipfs_cid: str = ""
    hash: str = ""

class SuspectAddRequest(BaseModel):
    case_id: str
    name: str
    father_name: str
    cnic: Optional[str] = None
    address: str
    description: str
    photo_cid: Optional[str] = None

class WitnessAddRequest(BaseModel):
    case_id: str
    name: str
    contact: str
    address: str
    statement: str

class EvidenceCreateRequest(BaseModel):
    case_id: str
    title: str
    description: str
    cloudinary_url: str
    file_hash: str

class EvidenceVerifyRequest(BaseModel):
    evidence_id: str
    provided_hash: str

class SubmitToCourtRequest(BaseModel):
    court_email: str

class SubmitToForensicRequest(BaseModel):
    forensic_email: str


# ============ CREATE CASE ============
@router.post("/", response_model=CaseResponse)
async def create_case(
    payload: CaseCreateRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    fir = await mongo_storage.get_fir(payload.fir_id)
    if not fir:
        raise HTTPException(status_code=404, detail="FIR not found")
    
    case_id = f"CASE-{uuid.uuid4().hex[:10].upper()}"
    case_number = f"JUSTICE-CASE-{datetime.now().year}-{uuid.uuid4().hex[:6].upper()}"
    
    # Get full user details for investigator name
    investigator_user = await mongo_storage.get_user(current_user["email"])
    investigator_name = investigator_user.get("full_name", current_user.get("full_name", current_user["email"]))
    
    case_data = {
        "case_id": case_id,
        "case_number": case_number,
        "fir_id": payload.fir_id,
        "title": payload.title,
        "description": payload.description,
        "priority": payload.priority,
        "status": "UNDER_INVESTIGATION",
        "investigator_email": current_user["email"],
        "investigator_name": investigator_name,
        "investigator_phone": investigator_user.get("phone_number", ""),
        "suspects": [],
        "witnesses": [],
        "evidence": [],
        "timeline": [{
            "action": "Case created",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "by": current_user["email"],
            "by_name": investigator_name
        }],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    # Generate hash for case
    case_hash = hashlib.sha256(json.dumps(case_data, default=str).encode()).hexdigest()
    case_data["hash"] = case_hash
    case_data["ipfs_cid"] = "STORED_IN_MONGODB"
    
    await mongo_storage.save_case(case_id, case_data)
    
    fir["status"] = "ACCEPTED"
    fir["assigned_investigator"] = current_user["email"]
    fir["case_id"] = case_id
    await mongo_storage.update_fir(payload.fir_id, fir)
    
    return CaseResponse(
        case_id=case_id,
        case_number=case_number,
        fir_id=payload.fir_id,
        title=payload.title,
        description=payload.description,
        status=case_data["status"],
        priority=payload.priority,
        investigator_email=current_user["email"],
        created_at=case_data["created_at"],
        updated_at=case_data["updated_at"],
        ipfs_cid=case_data.get("ipfs_cid", ""),
        hash=case_data.get("hash", "")
    )


# ============ ADD EVIDENCE ============
@router.post("/evidence")
async def add_evidence(
    payload: EvidenceCreateRequest,  # This will now accept cloudinary_url
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(payload.case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    evidence_id = f"EVD-{uuid.uuid4().hex[:10].upper()}"
    
    evidence_data = {
        "evidence_id": evidence_id,
        "case_id": payload.case_id,
        "title": payload.title,
        "description": payload.description,
        "cloudinary_url": payload.cloudinary_url,  # Changed from ipfs_cid
        "hash": payload.file_hash,
        "created_by": current_user["email"],
        "status": "COLLECTED",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "updated_at": datetime.now(timezone.utc).isoformat()
    }
    
    await mongo_storage.save_evidence(evidence_id, evidence_data)
    
    if "evidence" not in case:
        case["evidence"] = []
    case["evidence"].append(evidence_id)
    case["updated_at"] = datetime.now(timezone.utc).isoformat()
    await mongo_storage.update_case(payload.case_id, case)
    
    return {
        "evidence_id": evidence_id,
        "hash": payload.file_hash,
        "cloudinary_url": payload.cloudinary_url,  # Changed
        "title": payload.title,
        "case_id": payload.case_id,
        "status": "COLLECTED",
        "created_at": evidence_data["created_at"],
        "message": "Evidence added successfully"
    }


# ============ UPLOAD EVIDENCE FILE (CLOUDINARY) ============
@router.post("/evidence/upload-file")
async def add_evidence_upload_file(
    case_id: str = Form(...),
    title: str = Form(...),
    description: str = Form(""),
    file: UploadFile = File(...),
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR)),
):
    """Upload evidence file to Cloudinary and register metadata"""
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:  # 50MB limit
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB")
    file_hash = hashlib.sha256(content).hexdigest()
    
    # Upload to Cloudinary (instead of IPFS)
    cloudinary_result = await cloudinary_service.upload_file(
        content, 
        file.filename or "evidence.bin",
        folder=f"evidence/{case_id}"
    )
    
    if not cloudinary_result.get("url"):
        raise HTTPException(status_code=500, detail="Cloudinary upload failed")
    
    evidence_id = f"EVD-{uuid.uuid4().hex[:10].upper()}"
    now = datetime.now(timezone.utc).isoformat()
    
    evidence_data = {
        "evidence_id": evidence_id,
        "case_id": case_id,
        "title": title,
        "description": description,
        "cloudinary_url": cloudinary_result["url"],
        "cloudinary_public_id": cloudinary_result["public_id"],
        "hash": file_hash,
        "file_size": len(content),
        "mime_type": file.content_type or "application/octet-stream",
        "original_filename": file.filename or "evidence.bin",
        "created_by": current_user["email"],
        "status": "COLLECTED",
        "created_at": now,
        "updated_at": now,
    }

    await mongo_storage.save_evidence(evidence_id, evidence_data)
    
    if "evidence" not in case:
        case["evidence"] = []
    case["evidence"].append(evidence_id)
    case["updated_at"] = now
    await mongo_storage.update_case(case_id, case)

    return {
        "evidence_id": evidence_id,
        "hash": file_hash,
        "cloudinary_url": cloudinary_result["url"],
        "title": title,
        "case_id": case_id,
        "status": "COLLECTED",
        "created_at": evidence_data["created_at"],
        "message": "Evidence uploaded to Cloudinary successfully",
    }


# ============ VERIFY EVIDENCE BY FILE ============
@router.post("/evidence/verify-by-file")
async def verify_evidence_by_file(
    evidence_id: str,
    file: UploadFile = File(...),
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR, UserRole.FORENSIC_ANALYST, UserRole.COURT))
):
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    stored_hash = evidence.get("hash")
    stored_title = evidence.get("title", "N/A")
    
    file_content = await file.read()
    uploaded_file_name = file.filename
    uploaded_hash = hashlib.sha256(file_content).hexdigest()
    is_valid = uploaded_hash == stored_hash
    
    verification_record = {
        "verified_at": datetime.now(timezone.utc).isoformat(),
        "verified_by": current_user["email"],
        "uploaded_file_name": uploaded_file_name,
        "uploaded_hash": uploaded_hash,
        "stored_hash": stored_hash,
        "result": is_valid
    }
    
    if "verifications" not in evidence:
        evidence["verifications"] = []
    evidence["verifications"].append(verification_record)
    await mongo_storage.save_evidence(evidence_id, evidence)
    
    return {
        "evidence_id": evidence_id,
        "evidence_title": stored_title,
        "verified": is_valid,
        "uploaded_file_name": uploaded_file_name,
        "uploaded_hash": uploaded_hash,
        "stored_hash": stored_hash,
        "message": "✅ FILE IS AUTHENTIC!" if is_valid else "❌ FILE IS TAMPERED!",
        "verification_time": verification_record["verified_at"]
    }


# ============ GET INVESTIGATOR EVIDENCE ============
@router.get("/evidence/investigator")
async def get_investigator_evidence(current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))):
    all_cases = await mongo_storage.get_all_cases()
    investigator_case_ids = [case["case_id"] for case in all_cases if case.get("investigator_email") == current_user["email"]]
    all_evidence = await mongo_storage.get_all_evidence()
    evidence = [e for e in all_evidence if e.get("case_id") in investigator_case_ids]
    evidence.sort(key=lambda x: x.get("created_at", ""), reverse=True)
    return evidence


# ============ GET CASE EVIDENCE LIST ============
@router.get("/evidence/case/{case_id}")
async def get_case_evidence_list(
    case_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if case.get("investigator_email") != current_user["email"]:
        raise HTTPException(status_code=403, detail="Not your case")
    
    evidence_ids = case.get("evidence", [])
    evidence_list = []
    for ev_id in evidence_ids:
        ev = await mongo_storage.get_evidence(ev_id)
        if ev:
            evidence_list.append({
                "evidence_id": ev["evidence_id"],
                "title": ev["title"],
                "hash": ev.get("hash", ""),
                "ipfs_cid": ev.get("ipfs_cid", ""),
                "status": ev.get("status", "UNKNOWN"),
                "created_at": ev.get("created_at", "")
            })
    return evidence_list


# ============ GET EVIDENCE BY ID ============
@router.get("/evidence/{evidence_id}")
async def get_evidence(
    evidence_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR, UserRole.FORENSIC_ANALYST, UserRole.COURT))
):
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    return evidence


# ============ VERIFY EVIDENCE BY HASH ============
@router.post("/evidence/verify")
async def verify_evidence(
    payload: EvidenceVerifyRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR, UserRole.FORENSIC_ANALYST, UserRole.COURT))
):
    evidence = await mongo_storage.get_evidence(payload.evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    stored_hash = evidence.get("hash")
    is_valid = payload.provided_hash == stored_hash
    
    verification_record = {
        "verified_at": datetime.now(timezone.utc).isoformat(),
        "verified_by": current_user["email"],
        "provided_hash": payload.provided_hash,
        "stored_hash": stored_hash,
        "result": is_valid
    }
    
    if "verifications" not in evidence:
        evidence["verifications"] = []
    evidence["verifications"].append(verification_record)
    await mongo_storage.save_evidence(payload.evidence_id, evidence)
    
    return {
        "evidence_id": payload.evidence_id,
        "verified": is_valid,
        "stored_hash": stored_hash,
        "provided_hash": payload.provided_hash,
        "message": "✅ Evidence is AUTHENTIC!" if is_valid else "❌ Evidence TAMPERED!",
        "verification_time": verification_record["verified_at"]
    }


# ============ ADD SUSPECT ============
@router.post("/suspects")
async def add_suspect(
    payload: SuspectAddRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(payload.case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    suspect = {
        "id": str(uuid.uuid4()),
        "name": payload.name,
        "father_name": payload.father_name,
        "cnic": payload.cnic,
        "address": payload.address,
        "description": payload.description,
        "photo_cid": payload.photo_cid,
        "added_by": current_user["email"],
        "added_at": datetime.now(timezone.utc).isoformat()
    }
    
    case["suspects"].append(suspect)
    case["updated_at"] = datetime.now(timezone.utc).isoformat()
    await mongo_storage.update_case(payload.case_id, case)
    
    return {"status": "success", "suspect": suspect}


# ============ ADD WITNESS ============
@router.post("/witnesses")
async def add_witness(
    payload: WitnessAddRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(payload.case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    witness = {
        "id": str(uuid.uuid4()),
        "name": payload.name,
        "contact": payload.contact,
        "address": payload.address,
        "statement": payload.statement,
        "added_by": current_user["email"],
        "added_at": datetime.now(timezone.utc).isoformat()
    }
    
    case["witnesses"].append(witness)
    case["updated_at"] = datetime.now(timezone.utc).isoformat()
    await mongo_storage.update_case(payload.case_id, case)
    
    return {"status": "success", "witness": witness}


# ============ GET INVESTIGATOR CASES ============
@router.get("/investigator")
async def get_investigator_cases(current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))):
    all_cases = await mongo_storage.get_all_cases()
    cases = [c for c in all_cases if c.get("investigator_email") == current_user["email"]]
    return sorted(cases, key=lambda x: x.get("created_at", ""), reverse=True)


# ============ TRANSFER TO FORENSIC ============
@router.post("/evidence/transfer/forensic/{evidence_id}")
async def transfer_to_forensic(
    evidence_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    evidence["status"] = "TRANSFERRED_TO_FORENSIC"
    evidence["transferred_at"] = datetime.now(timezone.utc).isoformat()
    evidence["transferred_by"] = current_user["email"]
    await mongo_storage.save_evidence(evidence_id, evidence)
    
    return {"status": "ok", "evidence_id": evidence_id}


# ============ TRANSFER TO COURT ============
@router.post("/evidence/transfer/court/{evidence_id}")
async def transfer_to_court(
    evidence_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    evidence["status"] = "SUBMITTED_TO_COURT"
    evidence["transferred_to_court_at"] = datetime.now(timezone.utc).isoformat()
    evidence["transferred_by"] = current_user["email"]
    await mongo_storage.save_evidence(evidence_id, evidence)
    
    case = await mongo_storage.get_case(evidence["case_id"])
    if case:
        case["status"] = "UNDER_COURT_REVIEW"
        await mongo_storage.update_case(evidence["case_id"], case)
    
    return {"status": "ok", "evidence_id": evidence_id}


# ============ GET CASE BY ID ============
@router.get("/{case_id}")
async def get_case(case_id: str, current_user: dict = Depends(get_current_user)):
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return case


# ============ GET CASE EVIDENCE ============
@router.get("/{case_id}/evidence")
async def get_case_evidence(
    case_id: str,
    current_user: dict = Depends(get_current_user)
):
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    evidence_ids = case.get("evidence", [])
    evidence_list = []
    for ev_id in evidence_ids:
        ev = await mongo_storage.get_evidence(ev_id)
        if ev:
            evidence_list.append(ev)
    
    return evidence_list


# ============ GENERATE INVESTIGATION REPORT ============
@router.get("/generate-report/{case_id}")
async def generate_investigation_report(
    case_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    evidence_list = []
    for ev_id in case.get("evidence", []):
        ev = await mongo_storage.get_evidence(ev_id)
        if ev:
            evidence_list.append(ev)
    
    fir = await mongo_storage.get_fir(case.get("fir_id"))
    
    report = {
        "report_id": f"REP-{uuid.uuid4().hex[:8].upper()}",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "generated_by": current_user["email"],
        "case_details": {
            "case_id": case.get("case_id"),
            "case_number": case.get("case_number"),
            "title": case.get("title"),
            "description": case.get("description"),
            "status": case.get("status"),
            "priority": case.get("priority"),
            "created_at": case.get("created_at"),
            "investigator": case.get("investigator_email"),
            "investigator_name": case.get("investigator_name")
        },
        "fir_details": {
            "fir_id": fir.get("fir_id") if fir else None,
            "fir_number": fir.get("fir_number") if fir else None,
            "complainant_name": fir.get("complainant_name") if fir else None,
            "complainant_contact": fir.get("complainant_contact") if fir else None,
            "incident_title": fir.get("incident_title") if fir else None,
            "incident_description": fir.get("incident_description") if fir else None,
            "incident_location": fir.get("incident_location") if fir else None,
        },
        "suspects": case.get("suspects", []),
        "witnesses": case.get("witnesses", []),
        "evidence": [
            {
                "evidence_id": e.get("evidence_id"),
                "title": e.get("title"),
                "description": e.get("description"),
                "ipfs_cid": e.get("ipfs_cid"),
                "hash": e.get("hash"),
                "status": e.get("status"),
                "collected_by": e.get("created_by"),
                "collected_at": e.get("created_at"),
                "verifications": e.get("verifications", [])
            } for e in evidence_list
        ],
        "timeline": case.get("timeline", []),
        "evidence_verification_status": {
            "total": len(evidence_list),
            "verified": len([e for e in evidence_list if e.get("verifications")]),
            "pending": len([e for e in evidence_list if not e.get("verifications")]),
            "tampered": len([e for e in evidence_list if e.get("verifications") and not e.get("verifications")[-1].get("result")])
        },
        "status_summary": {
            "case_status": case.get("status"),
            "evidence_collected": len(evidence_list),
            "suspects_identified": len(case.get("suspects", [])),
            "witnesses_recorded": len(case.get("witnesses", []))
        }
    }
    
    await mongo_storage.save_report(report["report_id"], report)
    return report


# ============ DOWNLOAD REPORT ============
@router.get("/download-report/{case_id}")
async def download_investigation_report(
    case_id: str,
    format: str = "pdf",
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR, UserRole.COURT))
):
    """Download investigation report in PDF, DOCX, or TXT format"""
    
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    evidence_list = []
    for ev_id in case.get("evidence", []):
        ev = await mongo_storage.get_evidence(ev_id)
        if ev:
            evidence_list.append(ev)
    
    fir = await mongo_storage.get_fir(case.get("fir_id"))
    
    report_data = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "generated_by": current_user["email"],
        "case_details": {
            "case_id": case.get("case_id"),
            "case_number": case.get("case_number"),
            "title": case.get("title"),
            "description": case.get("description"),
            "status": case.get("status"),
            "priority": case.get("priority"),
            "created_at": case.get("created_at"),
            "investigator": case.get("investigator_email")
        },
        "fir_details": {
            "fir_number": fir.get("fir_number") if fir else None,
            "complainant_name": fir.get("complainant_name") if fir else None,
            "incident_title": fir.get("incident_title") if fir else None,
            "incident_description": fir.get("incident_description") if fir else None,
            "incident_location": fir.get("incident_location") if fir else None,
        },
        "suspects": case.get("suspects", []),
        "witnesses": case.get("witnesses", []),
        "evidence": [
            {
                "title": e.get("title"),
                "description": e.get("description", ""),
                "ipfs_cid": e.get("ipfs_cid"),
                "hash": e.get("hash"),
                "status": e.get("status")
            } for e in evidence_list
        ],
        "timeline": case.get("timeline", [])
    }
    
    case_number = case.get("case_number", case_id)
    
    if format.lower() == "pdf":
        return await generate_pdf_report(report_data, case_number)
    elif format.lower() == "docx":
        return await generate_docx_report(report_data, case_number)
    else:
        return await generate_txt_report(report_data, case_number)


async def generate_pdf_report(data: dict, case_number: str) -> Response:
    if not REPORTLAB_AVAILABLE:
        raise HTTPException(status_code=500, detail="PDF generation not available")
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    story = []
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#1e293b'), alignment=1, spaceAfter=30)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor('#3b82f6'), spaceBefore=20, spaceAfter=10)
    
    story.append(Paragraph(f"INVESTIGATION REPORT", title_style))
    story.append(Paragraph(f"Case Number: {case_number}", styles['Normal']))
    story.append(Paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("1. CASE DETAILS", heading_style))
    case_table = Table([
        ["Case Number", data['case_details'].get('case_number', 'N/A')],
        ["Title", data['case_details'].get('title', 'N/A')],
        ["Status", data['case_details'].get('status', 'N/A')],
        ["Priority", data['case_details'].get('priority', 'N/A')],
        ["Investigator", data['case_details'].get('investigator', 'N/A')]
    ], colWidths=[2*inch, 4*inch])
    case_table.setStyle(TableStyle([('BACKGROUND', (0,0), (0,-1), colors.lightgrey), ('GRID', (0,0), (-1,-1), 0.5, colors.grey)]))
    story.append(case_table)
    story.append(Spacer(1, 15))
    
    if data['fir_details'] and data['fir_details'].get('fir_number'):
        story.append(Paragraph("2. FIR DETAILS", heading_style))
        story.append(Paragraph(f"FIR Number: {data['fir_details'].get('fir_number', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"Complainant: {data['fir_details'].get('complainant_name', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"Incident: {data['fir_details'].get('incident_title', 'N/A')}", styles['Normal']))
        story.append(Paragraph(f"Location: {data['fir_details'].get('incident_location', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 15))
    
    if data['suspects']:
        story.append(Paragraph("3. SUSPECTS", heading_style))
        for s in data['suspects']:
            story.append(Paragraph(f"• {s.get('name', 'Unknown')}: {s.get('description', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    if data['witnesses']:
        story.append(Paragraph("4. WITNESSES", heading_style))
        for w in data['witnesses']:
            story.append(Paragraph(f"• {w.get('name', 'Unknown')}: {w.get('statement', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    if data['evidence']:
        story.append(Paragraph("5. EVIDENCE", heading_style))
        for e in data['evidence']:
            story.append(Paragraph(f"• {e.get('title', 'Unknown')} - Status: {e.get('status', 'N/A')}", styles['Normal']))
            story.append(Paragraph(f"  CID: {e.get('ipfs_cid', 'N/A')}", styles['Normal']))
        story.append(Spacer(1, 10))
    
    doc.build(story)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(), 
        media_type="application/pdf", 
        headers={
            "Access-Control-Allow-Origin": "*",
            "Content-Disposition": f"attachment; filename=investigation_report_{case_number}.pdf"
        }
    )


async def generate_docx_report(data: dict, case_number: str) -> Response:
    if not DOCX_AVAILABLE:
        raise HTTPException(status_code=500, detail="DOCX generation not available")
    
    doc = docx.Document()
    title = doc.add_heading('INVESTIGATION REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Case Number: {case_number}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()
    
    doc.add_heading('1. CASE DETAILS', level=1)
    doc.add_paragraph(f"Case Number: {data['case_details'].get('case_number', 'N/A')}")
    doc.add_paragraph(f"Title: {data['case_details'].get('title', 'N/A')}")
    doc.add_paragraph(f"Status: {data['case_details'].get('status', 'N/A')}")
    doc.add_paragraph(f"Investigator: {data['case_details'].get('investigator', 'N/A')}")
    doc.add_paragraph()
    
    if data['fir_details'] and data['fir_details'].get('fir_number'):
        doc.add_heading('2. FIR DETAILS', level=1)
        doc.add_paragraph(f"FIR Number: {data['fir_details'].get('fir_number', 'N/A')}")
        doc.add_paragraph(f"Complainant: {data['fir_details'].get('complainant_name', 'N/A')}")
        doc.add_paragraph(f"Incident: {data['fir_details'].get('incident_title', 'N/A')}")
        doc.add_paragraph()
    
    if data['suspects']:
        doc.add_heading('3. SUSPECTS', level=1)
        for s in data['suspects']:
            doc.add_paragraph(f"• {s.get('name', 'Unknown')}: {s.get('description', 'N/A')}")
        doc.add_paragraph()
    
    if data['witnesses']:
        doc.add_heading('4. WITNESSES', level=1)
        for w in data['witnesses']:
            doc.add_paragraph(f"• {w.get('name', 'Unknown')}: {w.get('statement', 'N/A')}")
        doc.add_paragraph()
    
    if data['evidence']:
        doc.add_heading('5. EVIDENCE', level=1)
        for e in data['evidence']:
            doc.add_paragraph(f"• {e.get('title', 'Unknown')}", style='List Bullet')
            doc.add_paragraph(f"  Status: {e.get('status', 'N/A')}")
            doc.add_paragraph(f"  IPFS CID: {e.get('ipfs_cid', 'N/A')}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(content=buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=investigation_report_{case_number}.docx"})


async def generate_txt_report(data: dict, case_number: str) -> Response:
    lines = []
    lines.append("=" * 80)
    lines.append("INVESTIGATION REPORT".center(80))
    lines.append("=" * 80)
    lines.append(f"Case Number: {case_number}")
    lines.append(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    
    lines.append("-" * 80)
    lines.append("1. CASE DETAILS")
    lines.append("-" * 80)
    lines.append(f"Case Number: {data['case_details'].get('case_number', 'N/A')}")
    lines.append(f"Title: {data['case_details'].get('title', 'N/A')}")
    lines.append(f"Status: {data['case_details'].get('status', 'N/A')}")
    lines.append(f"Investigator: {data['case_details'].get('investigator', 'N/A')}")
    lines.append("")
    
    if data['fir_details'] and data['fir_details'].get('fir_number'):
        lines.append("-" * 80)
        lines.append("2. FIR DETAILS")
        lines.append("-" * 80)
        lines.append(f"FIR Number: {data['fir_details'].get('fir_number', 'N/A')}")
        lines.append(f"Complainant: {data['fir_details'].get('complainant_name', 'N/A')}")
        lines.append(f"Incident: {data['fir_details'].get('incident_title', 'N/A')}")
        lines.append("")
    
    if data['suspects']:
        lines.append("-" * 80)
        lines.append("3. SUSPECTS")
        lines.append("-" * 80)
        for s in data['suspects']:
            lines.append(f"• {s.get('name', 'Unknown')}: {s.get('description', 'N/A')}")
        lines.append("")
    
    if data['witnesses']:
        lines.append("-" * 80)
        lines.append("4. WITNESSES")
        lines.append("-" * 80)
        for w in data['witnesses']:
            lines.append(f"• {w.get('name', 'Unknown')}: {w.get('statement', 'N/A')}")
        lines.append("")
    
    if data['evidence']:
        lines.append("-" * 80)
        lines.append("5. EVIDENCE")
        lines.append("-" * 80)
        for e in data['evidence']:
            lines.append(f"• {e.get('title', 'Unknown')}")
            lines.append(f"  Status: {e.get('status', 'N/A')}")
            lines.append(f"  IPFS CID: {e.get('ipfs_cid', 'N/A')}")
        lines.append("")
    
    lines.append("=" * 80)
    lines.append("END OF REPORT".center(80))
    lines.append("=" * 80)
    
    return Response(content="\n".join(lines).encode('utf-8'), media_type="text/plain", headers={"Content-Disposition": f"attachment; filename=investigation_report_{case_number}.txt"})


# ============ SUBMIT CASE TO COURT ============
@router.post("/submit-to-court/{case_id}")
async def submit_case_to_court(
    case_id: str,
    payload: SubmitToCourtRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    target_user = await mongo_storage.get_user(payload.court_email)
    if not target_user or target_user.get("role") != "COURT":
        raise HTTPException(status_code=404, detail="Court officer not found")
    
    for ev_id in case.get("evidence", []):
        ev = await mongo_storage.get_evidence(ev_id)
        if ev:
            ev["status"] = "SUBMITTED_TO_COURT"
            ev["submitted_to_court_at"] = datetime.now(timezone.utc).isoformat()
            ev["submitted_by"] = current_user["email"]
            ev["assigned_court_officer"] = payload.court_email
            ev["assigned_court_officer_name"] = target_user.get("full_name")
            await mongo_storage.save_evidence(ev_id, ev)
    
    case["status"] = "SUBMITTED_TO_COURT"
    case["submitted_to_court_at"] = datetime.now(timezone.utc).isoformat()
    case["submitted_by"] = current_user["email"]
    case["assigned_court_officer"] = payload.court_email
    case["assigned_court_officer_name"] = target_user.get("full_name")
    await mongo_storage.update_case(case_id, case)
    
    return {
        "message": f"Case submitted to {payload.court_email}", 
        "case_id": case_id, 
        "assigned_to": payload.court_email,
        "assigned_to_name": target_user.get("full_name")
    }


# ============ SUBMIT EVIDENCE TO FORENSIC ============
@router.post("/submit-to-forensic/{evidence_id}")
async def submit_evidence_to_forensic(
    evidence_id: str,
    payload: SubmitToForensicRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    target_user = await mongo_storage.get_user(payload.forensic_email)
    if not target_user or target_user.get("role") != "FORENSIC_ANALYST":
        raise HTTPException(status_code=404, detail="Forensic analyst not found")
    
    evidence["status"] = "TRANSFERRED_TO_FORENSIC"
    evidence["submitted_to_forensic_at"] = datetime.now(timezone.utc).isoformat()
    evidence["submitted_by"] = current_user["email"]
    evidence["assigned_forensic_analyst"] = payload.forensic_email
    evidence["assigned_forensic_analyst_name"] = target_user.get("full_name")
    evidence["analysis_status"] = "PENDING"
    await mongo_storage.save_evidence(evidence_id, evidence)
    
    case = await mongo_storage.get_case(evidence.get("case_id"))
    if case:
        if "timeline" not in case:
            case["timeline"] = []
        case["timeline"].append({
            "action": f"Evidence submitted to forensic lab",
            "timestamp": datetime.now(timezone.utc).isoformat(),
            "by": current_user["email"],
            "assigned_to": payload.forensic_email,
            "assigned_to_name": target_user.get("full_name")
        })
        await mongo_storage.update_case(evidence.get("case_id"), case)
    
    return {
        "message": f"Evidence submitted to {payload.forensic_email}", 
        "evidence_id": evidence_id, 
        "assigned_to": payload.forensic_email,
        "assigned_to_name": target_user.get("full_name")
    }


# ============ GET ALL CASES LIST ============
@router.get("/list")
async def get_all_cases_list(current_user: dict = Depends(get_current_user)):
    all_cases = await mongo_storage.get_all_cases()
    
    if current_user["role"] == UserRole.FORENSIC_ANALYST.value:
        cases_list = []
        all_evidence = await mongo_storage.get_all_evidence()
        for case in all_cases:
            case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case.get("case_id")]
            if case_evidence:
                cases_list.append({
                    "case_id": case.get("case_id"),
                    "case_number": case.get("case_number"),
                    "title": case.get("title"),
                    "status": case.get("status")
                })
        return cases_list
    else:
        return [
            {
                "case_id": case.get("case_id"),
                "case_number": case.get("case_number"),
                "title": case.get("title"),
                "status": case.get("status")
            }
            for case in all_cases
        ]


# ============ GET MY CASES ============
@router.get("/my-cases")
async def get_my_cases(current_user: dict = Depends(get_current_user)):
    all_cases = await mongo_storage.get_all_cases()
    
    my_cases = []
    for case in all_cases:
        if case.get("investigator_email") == current_user["email"]:
            my_cases.append({
                "case_id": case.get("case_id"),
                "case_number": case.get("case_number"),
                "title": case.get("title"),
                "status": case.get("status"),
                "created_at": case.get("created_at")
            })
    
    return my_cases


# ============ SEND SINGLE EVIDENCE TO FORENSIC ============
@router.post("/{case_id}/evidence/{evidence_id}/send-to-forensic")
async def send_single_evidence_to_forensic(
    case_id: str,
    evidence_id: str,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    case = await mongo_storage.get_case(case_id)
    if not case or case.get("investigator_email") != current_user["email"]:
        raise HTTPException(status_code=403, detail="Access denied")
    
    evidence = await mongo_storage.get_evidence(evidence_id)
    if not evidence or evidence.get("case_id") != case_id:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    evidence["status"] = "TRANSFERRED_TO_FORENSIC"
    evidence["analysis_status"] = "PENDING"
    evidence["transferred_at"] = datetime.now(timezone.utc).isoformat()
    evidence["transferred_by"] = current_user["email"]
    evidence["transferred_by_name"] = current_user.get("full_name")
    
    await mongo_storage.save_evidence(evidence_id, evidence)
    
    if case.get("forensic_status") != "ACCEPTED":
        case["forensic_status"] = "PENDING"
        await mongo_storage.save_case(case_id, case)
    
    return {
        "message": f"Evidence '{evidence.get('title')}' sent to forensic lab",
        "evidence_id": evidence_id,
        "case_id": case_id
    }


# ============ GET MY HEARINGS ============
@router.get("/my-hearings")
async def get_my_hearings(current_user: dict = Depends(get_current_user)):
    all_cases = await mongo_storage.get_all_cases()
    my_hearings = []
    
    user_email = current_user["email"]
    user_role = current_user["role"]
    
    for case in all_cases:
        fir = await mongo_storage.get_fir(case.get("fir_id"))
        
        is_complainant = fir and fir.get("complainant_email") == user_email
        is_investigator = case.get("investigator_email") == user_email
        
        if is_complainant or is_investigator or user_role == "ADMIN":
            hearings = await mongo_storage.get_case_hearings(case.get("case_id"))
            for hearing_id, hearing in hearings.items():
                my_hearings.append({
                    "hearing_id": hearing_id,
                    "case_id": case.get("case_id"),
                    "case_number": case.get("case_number"),
                    "case_title": case.get("title"),
                    "hearing_date": hearing.get("hearing_date"),
                    "hearing_time": hearing.get("hearing_time"),
                    "hearing_type": hearing.get("hearing_type"),
                    "meeting_link": hearing.get("meeting_link"),
                    "status": hearing.get("status"),
                    "notes": hearing.get("notes"),
                    "scheduled_by_name": hearing.get("scheduled_by_name"),
                    "scheduled_by": hearing.get("scheduled_by"),
                    "created_at": hearing.get("created_at")
                })
    
    my_hearings.sort(key=lambda x: x.get("hearing_date", ""))
    return my_hearings