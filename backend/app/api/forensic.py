# backend/app/api/forensic.py
import hashlib
import json
from fastapi import APIRouter, Depends, HTTPException, status
from pydantic import BaseModel
from datetime import datetime, timezone
from typing import Optional, List
import uuid
from app.core.authz import require_roles
from app.core.roles import UserRole
from app.core.authz import get_current_user
from app.services.mongo_storage import mongo_storage
from fastapi.responses import FileResponse
import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import Response
from app.api.websocket import manager

router = APIRouter(prefix="/forensic", tags=["Forensic"])

class EvidenceTransferRequest(BaseModel):
    evidence_id: str
    case_id: str

class ForensicAnalysisRequest(BaseModel):
    evidence_id: str
    analysis_type: str
    findings: str
    conclusion: str
    methodology: Optional[str] = None
    equipment_used: Optional[str] = None

class ForensicReportResponse(BaseModel):
    report_id: str
    report_number: str
    evidence_id: str
    case_id: str
    analysis_type: str
    findings: str
    conclusion: str
    analyst_email: str
    status: str
    storage_cid: str
    hash: str
    created_at: str

class TransferToCourtRequest(BaseModel):
    court_email: str

@router.post("/transfer")
async def transfer_to_forensic(
    payload: EvidenceTransferRequest,
    current_user: dict = Depends(require_roles(UserRole.INVESTIGATOR))
):
    evidence = await mongo_storage.get_evidence(payload.evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    evidence["status"] = "TRANSFERRED_TO_FORENSIC"
    evidence["transferred_at"] = datetime.now(timezone.utc).isoformat()
    evidence["transferred_by"] = current_user["email"]
    evidence["analysis_status"] = "PENDING"
    
    await mongo_storage.save_evidence(payload.evidence_id, evidence)
    
    return {"status": "success", "message": "Evidence transferred to forensic lab"}

@router.get("/queue")
async def get_forensic_queue(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    all_evidence = await mongo_storage.get_all_evidence()
    
    queue = []
    for evidence in all_evidence:
        # Check multiple conditions for pending analysis
        is_pending = evidence.get("status") == "TRANSFERRED_TO_FORENSIC"
        
        if is_pending:
            queue.append({
                "evidence_id": evidence.get("evidence_id"),
                "case_id": evidence.get("case_id"),
                "title": evidence.get("title"),
                "description": evidence.get("description"),
                "cloudinary_url": evidence.get("cloudinary_url") or evidence.get("ipfs_cid"),
                "hash": evidence.get("hash"),
                "status": evidence.get("status"),
                "analysis_status": evidence.get("analysis_status", "PENDING"),
                "priority": evidence.get("priority", "MEDIUM"),
                "transferred_at": evidence.get("transferred_at") or evidence.get("created_at"),
                "transferred_by": evidence.get("transferred_by") or evidence.get("created_by")
            })
    
    # Sort with None handling
    return sorted(queue, key=lambda x: x.get("transferred_at") or "", reverse=True)

@router.post("/analyze", response_model=ForensicReportResponse)
async def create_forensic_report(
    payload: ForensicAnalysisRequest,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    evidence = await mongo_storage.get_evidence(payload.evidence_id)
    if not evidence:
        raise HTTPException(status_code=404, detail="Evidence not found")
    
    case = await mongo_storage.get_case(evidence.get("case_id"))
    
    report_id = f"FR-{uuid.uuid4().hex[:10].upper()}"
    report_number = f"FORENSIC-{datetime.now().year}-{uuid.uuid4().hex[:6].upper()}"
    
    report_data = {
        "report_id": report_id,
        "report_number": report_number,
        "evidence_id": payload.evidence_id,
        "case_id": evidence.get("case_id"),
        "case_title": case.get("title") if case else "Unknown",
        "case_number": case.get("case_number") if case else "Unknown",
        "analysis_type": payload.analysis_type,
        "findings": payload.findings,
        "conclusion": payload.conclusion,
        "methodology": getattr(payload, "methodology", ""),
        "equipment_used": getattr(payload, "equipment_used", ""),
        "analyst_email": current_user["email"],
        "analyst_name": current_user.get("full_name", ""),
        "status": "COMPLETED",
        "created_at": datetime.now(timezone.utc).isoformat()
    }
    
    report_hash = hashlib.sha256(json.dumps(report_data, default=str).encode()).hexdigest()
    report_data["hash"] = report_hash
    report_data["storage_cid"] = "STORED_IN_MONGODB"
    
    # Save report
    forensic_reports = await mongo_storage.get_forensic_reports()
    forensic_reports[report_id] = report_data
    await mongo_storage.save_forensic_reports(forensic_reports)
    
    # Update evidence
    evidence["forensic_report_id"] = report_id
    evidence["forensic_report_number"] = report_number
    evidence["analysis_status"] = "COMPLETED"
    evidence["analysis_completed_at"] = datetime.now(timezone.utc).isoformat()
    await mongo_storage.save_evidence(payload.evidence_id, evidence)
    
    return ForensicReportResponse(
        report_id=report_id,
        report_number=report_number,
        evidence_id=payload.evidence_id,
        case_id=evidence.get("case_id"),
        analysis_type=payload.analysis_type,
        findings=payload.findings,
        conclusion=payload.conclusion,
        analyst_email=current_user["email"],
        status="COMPLETED",
        storage_cid=report_data.get("storage_cid", ""),
        hash=report_data.get("hash", ""),
        created_at=report_data["created_at"]
    )

@router.post("/transfer-court/{report_id}")
async def transfer_report_to_court(
    report_id: str,
    payload: TransferToCourtRequest,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    forensic_reports = await mongo_storage.get_forensic_reports()
    report = forensic_reports.get(report_id)
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    target_user = await mongo_storage.get_user(payload.court_email)
    if not target_user or target_user.get("role") != "COURT":
        raise HTTPException(status_code=404, detail="Court officer not found")
    
    report["status"] = "TRANSFERRED_TO_COURT"
    report["transferred_at"] = datetime.now(timezone.utc).isoformat()
    report["transferred_by"] = current_user["email"]
    report["assigned_court_officer"] = payload.court_email
    report["assigned_court_officer_name"] = target_user.get("full_name")
    
    forensic_reports[report_id] = report
    await mongo_storage.save_forensic_reports(forensic_reports)
    
    evidence = await mongo_storage.get_evidence(report.get("evidence_id"))
    if evidence:
        evidence["status"] = "SUBMITTED_TO_COURT"
        evidence["transferred_to_court_at"] = datetime.now(timezone.utc).isoformat()
        evidence["assigned_court_officer"] = payload.court_email
        await mongo_storage.save_evidence(report.get("evidence_id"), evidence)
    
    return {
        "status": "success", 
        "message": f"Report transferred to court officer {payload.court_email}",
        "assigned_to": payload.court_email,
        "assigned_to_name": target_user.get("full_name")
    }

@router.get("/reports")
async def get_forensic_reports(current_user: dict = Depends(get_current_user)):
    forensic_reports = await mongo_storage.get_forensic_reports()
    reports_list = list(forensic_reports.values())
    
    if current_user["role"] in [UserRole.FORENSIC_ANALYST.value, UserRole.COURT.value]:
        return sorted(reports_list, key=lambda x: x.get("created_at", ""), reverse=True)
    
    if current_user["role"] == UserRole.INVESTIGATOR.value:
        all_cases = await mongo_storage.get_all_cases()
        my_cases = [c for c in all_cases if c.get("investigator_email") == current_user["email"]]
        my_case_ids = [c.get("case_id") for c in my_cases]
        
        filtered_reports = [r for r in reports_list if r.get("case_id") in my_case_ids]
        return sorted(filtered_reports, key=lambda x: x.get("created_at", ""), reverse=True)
    
    return []

@router.get("/report/{report_id}")
async def get_forensic_report(report_id: str, current_user: dict = Depends(get_current_user)):
    forensic_reports = await mongo_storage.get_forensic_reports()
    report = forensic_reports.get(report_id)
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    if current_user["role"] == UserRole.INVESTIGATOR.value:
        case = await mongo_storage.get_case(report.get("case_id"))
        if case and case.get("investigator_email") != current_user["email"]:
            raise HTTPException(status_code=403, detail="Access denied")
    
    return report

@router.get("/stats")
async def get_forensic_stats(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get forensic analyst statistics"""
    all_evidence = await mongo_storage.get_all_evidence()
    forensic_reports = await mongo_storage.get_forensic_reports()
    all_cases = await mongo_storage.get_all_cases()
    
    # CASE STATS - Based on case.forensic_status
    pending_cases = []
    accepted_cases = []
    for case in all_cases:
        if case.get("forensic_status") == "PENDING" or case.get("forensic_status") is None:
            # Check if case has any evidence sent to forensic
            case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case.get("case_id")]
            has_sent_to_forensic = any(
                ev.get("status") == "TRANSFERRED_TO_FORENSIC" or 
                ev.get("analysis_status") in ["PENDING", None]
                for ev in case_evidence
            )
            if has_sent_to_forensic:
                pending_cases.append(case)
        elif case.get("forensic_status") == "ACCEPTED":
            accepted_cases.append(case)
    
    # Evidence stats
    pending_evidence = len([e for e in all_evidence if e.get("status") == "TRANSFERRED_TO_FORENSIC" and e.get("analysis_status") in ["PENDING", None]])
    in_progress_evidence = len([e for e in all_evidence if e.get("analysis_status") == "IN_PROGRESS"])
    completed_evidence = len([e for e in all_evidence if e.get("analysis_status") == "COMPLETED"])
    
    my_reports = [r for r in forensic_reports.values() if r.get("analyst_email") == current_user["email"]]
    
    return {
        "pending_cases": len(pending_cases),
        "accepted_cases": len(accepted_cases),
        "total_cases": len(all_cases),
        "pending_evidence": pending_evidence,
        "in_progress_evidence": in_progress_evidence, 
        "completed_evidence": completed_evidence,
        "total_reports": len(forensic_reports.values()),
        "my_reports": len(my_reports),
        "transferred_to_court": len([r for r in forensic_reports.values() if r.get("status") == "TRANSFERRED_TO_COURT"])
    }

@router.get("/forensic-evidence")
async def get_forensic_evidence(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get all evidence for forensic analyst"""
    all_evidence = await mongo_storage.get_all_evidence()
    
    forensic_evidence = []
    for evidence in all_evidence:
        forensic_evidence.append({
            "evidence_id": evidence.get("evidence_id"),
            "case_id": evidence.get("case_id"),
            "title": evidence.get("title"),
            "description": evidence.get("description"),
            "cloudinary_url": evidence.get("cloudinary_url") or evidence.get("ipfs_cid"),
            "hash": evidence.get("hash"),
            "status": evidence.get("status"),
            "analysis_status": evidence.get("analysis_status", "PENDING"),
            "created_at": evidence.get("created_at")
        })
    
    return forensic_evidence

@router.get("/cases-list")
async def get_forensic_cases(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get all cases for forensic analyst with investigator info"""
    all_cases = await mongo_storage.get_all_cases()
    all_evidence = await mongo_storage.get_all_evidence()
    
    cases_list = []
    for case in all_cases:
        # Find who submitted evidence to forensic for this case
        investigator_email = None
        investigator_name = None
        
        case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case.get("case_id")]
        
        # Check if any evidence was sent to forensic
        for ev in case_evidence:
            if ev.get("transferred_by"):
                investigator_email = ev.get("transferred_by")
                investigator_name = ev.get("transferred_by_name")
                break
        
        # If no evidence sent, fallback to case investigator
        if not investigator_email:
            investigator_email = case.get("investigator_email")
            investigator_name = case.get("investigator_name")
        
        cases_list.append({
            "case_id": case.get("case_id"),
            "case_number": case.get("case_number"),
            "title": case.get("title"),
            "status": case.get("status"),
            "investigator_email": investigator_email,
            "investigator_name": investigator_name,
            "evidence_count": len(case_evidence),
            "has_forensic_submission": any(ev.get("transferred_at") for ev in case_evidence)
        })
    
    return cases_list

@router.get("/case-investigator/{case_id}")
async def get_case_investigator(
    case_id: str,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    """Get the investigator who submitted evidence for this case"""
    all_evidence = await mongo_storage.get_all_evidence()
    case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case_id]
    
    for ev in case_evidence:
        if ev.get("transferred_at") and ev.get("transferred_by"):
            return {
                "email": ev.get("transferred_by"),
                "name": ev.get("transferred_by_name") or ev.get("transferred_by"),
            }
    
    case = await mongo_storage.get_case(case_id)
    return {
        "email": case.get("investigator_email") if case else None,
        "name": case.get("investigator_name") if case else None
    }

@router.get("/reports/by-case/{case_id}")
async def get_reports_by_case(case_id: str, current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get reports for a specific case"""
    forensic_reports = await mongo_storage.get_forensic_reports()
    reports_list = list(forensic_reports.values())
    
    case_reports = [r for r in reports_list if r.get("case_id") == case_id]
    return sorted(case_reports, key=lambda x: x.get("created_at", ""), reverse=True)


class ForensicCaseAction(BaseModel):
    action: str  # ACCEPT, REJECT
    reason: Optional[str] = None
    estimated_days: Optional[int] = None

@router.get("/pending-cases")
async def get_pending_forensic_cases(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get all cases pending forensic acceptance"""
    all_cases = await mongo_storage.get_all_cases()
    all_evidence = await mongo_storage.get_all_evidence()
    
    pending_cases = []
    for case in all_cases:
        if case.get("forensic_status") == "ACCEPTED":
            continue
        
        case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case.get("case_id")]
        has_forensic_pending = any(
    ev.get("status") == "TRANSFERRED_TO_FORENSIC"
    for ev in case_evidence
)
        
        if has_forensic_pending:
            fir = await mongo_storage.get_fir(case.get("fir_id"))
            investigator = None
            if case.get("investigator_email"):
                investigator = await mongo_storage.get_user(case.get("investigator_email"))
            
            pending_cases.append({
                "case_id": case.get("case_id"),
                "case_number": case.get("case_number"),
                "title": case.get("title"),
                "description": case.get("description"),
                "fir": {
                    "fir_number": fir.get("fir_number") if fir else "N/A",
                    "complainant_name": fir.get("complainant_name") if fir else "N/A",
                    "complaint_text": fir.get("complaint_text") if fir else "N/A",
                    "filed_at": fir.get("created_at") if fir else None
                },
                "investigator": {
                    "name": investigator.get("full_name") if investigator else case.get("investigator_name", "Not Assigned"),
                    "email": case.get("investigator_email"),
                    "badge_number": investigator.get("badge_number") if investigator else "N/A"
                } if investigator else None,
                "evidence_list": [
                    {
                        "evidence_id": ev.get("evidence_id"),
                        "title": ev.get("title"),
                        "description": ev.get("description"),
                        "status": ev.get("status"),
                        "analysis_status": ev.get("analysis_status", "PENDING"),
                        "submitted_at": ev.get("transferred_at") or ev.get("created_at"),
                        "submitted_by": ev.get("transferred_by") or ev.get("created_by")
                    }
                    for ev in case_evidence
                    if ev.get("status") == "TRANSFERRED_TO_FORENSIC" or ev.get("analysis_status") in ["PENDING", None]
                ],
                "total_evidence": len(case_evidence),
                "submitted_at": min([ev.get("transferred_at") or ev.get("created_at") for ev in case_evidence if ev.get("transferred_at")], default=None),
                "priority": "HIGH" if any(ev.get("priority") == "HIGH" for ev in case_evidence) else "MEDIUM"
            })
    
    pending_cases.sort(key=lambda x: x.get("submitted_at") or "", reverse=True)
    return pending_cases

@router.post("/case-action/{case_id}")
async def forensic_case_action(
    case_id: str,
    payload: ForensicCaseAction,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    """Accept or reject ENTIRE CASE for forensic analysis"""
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    all_evidence = await mongo_storage.get_all_evidence()
    case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case_id]
    
    action_timestamp = datetime.now(timezone.utc).isoformat()
    
    if payload.action == "ACCEPT":
        for ev in case_evidence:
            if ev.get("status") == "TRANSFERRED_TO_FORENSIC" or ev.get("analysis_status") in ["PENDING", None]:
                ev["analysis_status"] = "ACCEPTED"
                ev["accepted_by"] = current_user["email"]
                ev["accepted_by_name"] = current_user.get("full_name")
                ev["accepted_at"] = action_timestamp
                ev["estimated_days"] = payload.estimated_days
                ev["case_accepted_at"] = action_timestamp
                await mongo_storage.save_evidence(ev.get("evidence_id"), ev)
        
        case["forensic_status"] = "ACCEPTED"
        case["forensic_accepted_at"] = action_timestamp
        case["forensic_accepted_by"] = current_user["email"]
        case["forensic_accepted_by_name"] = current_user.get("full_name")
        case["forensic_estimated_days"] = payload.estimated_days
        await mongo_storage.save_case(case_id, case)
        
        return {
            "status": "success",
            "message": f"Case {case_id} accepted for forensic analysis",
            "estimated_days": payload.estimated_days
        }
    
    elif payload.action == "REJECT":
        reason = payload.reason or "Not specified"
        
        for ev in case_evidence:
            if ev.get("status") == "TRANSFERRED_TO_FORENSIC" or ev.get("analysis_status") in ["PENDING", None]:
                ev["analysis_status"] = "REJECTED"
                ev["rejected_by"] = current_user["email"]
                ev["rejected_by_name"] = current_user.get("full_name")
                ev["rejected_at"] = action_timestamp
                ev["rejection_reason"] = reason
                ev["status"] = "REJECTED_BY_FORENSIC"
                await mongo_storage.save_evidence(ev.get("evidence_id"), ev)
        
        case["forensic_status"] = "REJECTED"
        case["forensic_rejected_at"] = action_timestamp
        case["forensic_rejected_by"] = current_user["email"]
        case["forensic_rejected_by_name"] = current_user.get("full_name")
        case["forensic_rejection_reason"] = reason
        await mongo_storage.save_case(case_id, case)
        
        return {
            "status": "success",
            "message": f"Case {case_id} rejected",
            "reason": reason
        }
    
    raise HTTPException(status_code=400, detail="Invalid action")

@router.get("/case-details/{case_id}")
async def get_full_case_for_forensic(
    case_id: str,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    """Get complete case details including FIR, investigator actions, evidence"""
    case = await mongo_storage.get_case(case_id)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    
    fir = await mongo_storage.get_fir(case.get("fir_id"))
    investigator = None
    if case.get("investigator_email"):
        investigator = await mongo_storage.get_user(case.get("investigator_email"))
    
    all_evidence = await mongo_storage.get_all_evidence()
    case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case_id]
    
    # Build complete timeline
    timeline = []
    
    # 1. FIR Submission
    if fir and fir.get("created_at"):
        timeline.append({
            "date": fir.get("created_at"),
            "event": "📝 FIR Submitted",
            "description": f"FIR #{fir.get('fir_number')} filed by {fir.get('complainant_name')}",
            "icon": "📝",
            "by": fir.get("complainant_name"),
            "details": {
                "fir_number": fir.get("fir_number"),
                "complaint_text": fir.get("complaint_text"),
                "incident_date": fir.get("incident_date"),
                "incident_location": fir.get("incident_location")
            }
        })
    
    # 2. FIR Status Changes
    if fir and fir.get("status_history"):
        for history in fir.get("status_history", []):
            timeline.append({
                "date": history.get("timestamp"),
                "event": f"📌 FIR {history.get('status', '').replace('_', ' ')}",
                "description": history.get("remarks", f"FIR status updated to {history.get('status')}"),
                "icon": "📌",
                "by": history.get("changed_by", "System")
            })
    
    # 3. Case Created
    if case.get("created_at"):
        timeline.append({
            "date": case.get("created_at"),
            "event": "⚖️ Case Created",
            "description": f"Case #{case.get('case_number')} registered and assigned to investigator",
            "icon": "⚖️",
            "by": case.get("investigator_name", "System")
        })
    
    # 4. Suspects Added
    for suspect in case.get("suspects", []):
        timeline.append({
            "date": suspect.get("added_at") or case.get("created_at"),
            "event": "👤 Suspect Identified",
            "description": f"Suspect: {suspect.get('name')} - {suspect.get('description', '')}",
            "icon": "👤",
            "by": suspect.get("added_by", "Investigator")
        })
    
    # 5. Witnesses Added
    for witness in case.get("witnesses", []):
        timeline.append({
            "date": witness.get("added_at") or case.get("created_at"),
            "event": "👥 Witness Recorded",
            "description": f"Witness: {witness.get('name')}",
            "icon": "👥",
            "by": witness.get("added_by", "Investigator")
        })
    
    # 6. Evidence Added and Sent to Forensic
    for ev in case_evidence:
        if ev.get("created_at"):
            timeline.append({
                "date": ev.get("created_at"),
                "event": "📦 Evidence Added",
                "description": f"Evidence: {ev.get('title')} - {ev.get('description', '')}",
                "icon": "📦",
                "by": ev.get("created_by", "Investigator"),
                "evidence_id": ev.get("evidence_id"),
                "evidence_title": ev.get("title")
            })
        
        if ev.get("transferred_at"):
            timeline.append({
                "date": ev.get("transferred_at"),
                "event": "🔬 Evidence Sent to Forensic Lab",
                "description": f"Evidence '{ev.get('title')}' submitted for forensic analysis",
                "icon": "🔬",
                "by": ev.get("transferred_by", "Investigator"),
                "evidence_id": ev.get("evidence_id")
            })
        
        if ev.get("accepted_at"):
            timeline.append({
                "date": ev.get("accepted_at"),
                "event": "✅ Evidence Accepted by Forensic",
                "description": f"Evidence '{ev.get('title')}' accepted for analysis. Estimated: {ev.get('estimated_days', 7)} days",
                "icon": "✅",
                "by": ev.get("accepted_by_name", "Forensic Analyst"),
                "evidence_id": ev.get("evidence_id")
            })
        
        if ev.get("rejected_at"):
            timeline.append({
                "date": ev.get("rejected_at"),
                "event": "❌ Evidence Rejected by Forensic",
                "description": f"Evidence '{ev.get('title')}' rejected. Reason: {ev.get('rejection_reason', 'Not specified')}",
                "icon": "❌",
                "by": ev.get("rejected_by_name", "Forensic Analyst"),
                "evidence_id": ev.get("evidence_id")
            })
        
        if ev.get("analysis_completed_at"):
            timeline.append({
                "date": ev.get("analysis_completed_at"),
                "event": "📊 Forensic Analysis Completed",
                "description": f"Analysis complete for '{ev.get('title')}'. Report available.",
                "icon": "📊",
                "by": ev.get("analyzed_by_name", "Forensic Analyst"),
                "evidence_id": ev.get("evidence_id")
            })
    
    # 7. Case status updates
    if case.get("submitted_to_court_at"):
        timeline.append({
            "date": case.get("submitted_to_court_at"),
            "event": "🏛️ Case Submitted to Court",
            "description": f"Case submitted to court officer",
            "icon": "🏛️",
            "by": case.get("submitted_by", "Investigator")
        })
    
    timeline.sort(key=lambda x: x.get("date", ""))
    
    forensic_reports = await mongo_storage.get_forensic_reports()
    
    evidence_with_details = []
    for ev in case_evidence:
        report = None
        for r_id, r in forensic_reports.items():
            if r.get("evidence_id") == ev.get("evidence_id"):
                report = r
                break
        
        evidence_with_details.append({
            "evidence_id": ev.get("evidence_id"),
            "title": ev.get("title"),
            "description": ev.get("description"),
            "type": ev.get("type", "DOCUMENT"),
            "hash": ev.get("hash"),
            "cloudinary_url": ev.get("cloudinary_url") or ev.get("ipfs_cid"),
            "status": ev.get("status"),
            "analysis_status": ev.get("analysis_status", "PENDING"),
            "submitted_to_forensic_at": ev.get("transferred_at"),
            "accepted_at": ev.get("accepted_at"),
            "rejected_at": ev.get("rejected_at"),
            "rejection_reason": ev.get("rejection_reason"),
            "analysis_completed_at": ev.get("analysis_completed_at"),
            "created_at": ev.get("created_at"),
            "created_by": ev.get("created_by"),
            "forensic_report": {
                "report_id": report.get("report_id") if report else None,
                "report_number": report.get("report_number") if report else None,
                "findings": report.get("findings") if report else None,
                "conclusion": report.get("conclusion") if report else None,
                "created_at": report.get("created_at") if report else None
            } if report else None
        })
    
    return {
        "case_id": case.get("case_id"),
        "case_number": case.get("case_number"),
        "title": case.get("title"),
        "description": case.get("description"),
        "status": case.get("status"),
        "forensic_status": case.get("forensic_status", "PENDING"),
        "fir": {
            "fir_number": fir.get("fir_number") if fir else "N/A",
            "filed_at": fir.get("created_at") if fir else None,
            "complainant": {
                "name": fir.get("complainant_name") if fir else "N/A",
                "email": fir.get("complainant_email") if fir else "N/A",
                "phone": fir.get("complainant_phone") if fir else "N/A"
            } if fir else None,
            "complaint_text": fir.get("complaint_text") if fir else "N/A",
            "incident_date": fir.get("incident_date") if fir else None,
            "incident_location": fir.get("incident_location") if fir else "N/A",
            "status_history": fir.get("status_history", []) if fir else []
        },
        "investigator": {
            "name": investigator.get("full_name") if investigator else case.get("investigator_name", "Not Assigned"),
            "email": case.get("investigator_email"),
            "badge_number": investigator.get("badge_number") if investigator else "N/A",
            "phone": investigator.get("phone") if investigator else "N/A"
        } if investigator or case.get("investigator_email") else None,
        "suspects": case.get("suspects", []),
        "witnesses": case.get("witnesses", []),
        "evidence": evidence_with_details,
        "timeline": timeline,
        "created_at": case.get("created_at"),
        "updated_at": case.get("updated_at")
    }

@router.get("/accepted-cases")
async def get_accepted_forensic_cases(current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))):
    """Get all cases accepted by forensic for analysis"""
    all_cases = await mongo_storage.get_all_cases()
    all_evidence = await mongo_storage.get_all_evidence()
    
    accepted_cases = []
    for case in all_cases:
        case_evidence = [ev for ev in all_evidence if ev.get("case_id") == case.get("case_id")]
        has_accepted = any(ev.get("analysis_status") == "ACCEPTED" or ev.get("accepted_at") for ev in case_evidence)
        has_completed = any(ev.get("analysis_status") == "COMPLETED" for ev in case_evidence)
        
        if has_accepted or has_completed:
            fir = await mongo_storage.get_fir(case.get("fir_id"))
            investigator = None
            if case.get("investigator_email"):
                investigator = await mongo_storage.get_user(case.get("investigator_email"))
            
            total_evidence = len(case_evidence)
            accepted_count = len([ev for ev in case_evidence if ev.get("analysis_status") == "ACCEPTED" or ev.get("accepted_at")])
            completed_count = len([ev for ev in case_evidence if ev.get("analysis_status") == "COMPLETED"])
            pending_count = len([ev for ev in case_evidence if ev.get("analysis_status") == "PENDING" or ev.get("analysis_status") is None])
            
            accepted_cases.append({
                "case_id": case.get("case_id"),
                "case_number": case.get("case_number"),
                "title": case.get("title"),
                "description": case.get("description"),
                "status": case.get("status"),
                "fir": {
                    "fir_number": fir.get("fir_number") if fir else "N/A",
                    "complainant_name": fir.get("complainant_name") if fir else "N/A",
                    "complaint_text": fir.get("complaint_text") if fir else "N/A",
                    "filed_at": fir.get("created_at") if fir else None
                },
                "investigator": {
                    "name": investigator.get("full_name") if investigator else case.get("investigator_name", "Not Assigned"),
                    "email": case.get("investigator_email"),
                    "badge_number": investigator.get("badge_number") if investigator else "N/A"
                } if investigator else None,
                "evidence_stats": {
                    "total": total_evidence,
                    "accepted": accepted_count,
                    "completed": completed_count,
                    "pending": pending_count
                },
                "accepted_at": min([ev.get("accepted_at") for ev in case_evidence if ev.get("accepted_at")], default=None),
                "priority": "HIGH" if any(ev.get("priority") == "HIGH" for ev in case_evidence) else "MEDIUM"
            })
    
    accepted_cases.sort(key=lambda x: x.get("accepted_at") or "", reverse=True)
    return accepted_cases

@router.get("/report/download/{report_id}")
async def download_forensic_report(
    report_id: str,
    format: str = "pdf",
    current_user: dict = Depends(get_current_user)
):
    """Download forensic report in PDF, DOCX, or TXT format"""
    forensic_reports = await mongo_storage.get_forensic_reports()
    report = forensic_reports.get(report_id)
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    # Check permissions
    case = await mongo_storage.get_case(report.get("case_id"))
    if current_user["role"] == UserRole.INVESTIGATOR.value:
        if case and case.get("investigator_email") != current_user["email"]:
            raise HTTPException(status_code=403, detail="Access denied")
    elif current_user["role"] == UserRole.PUBLIC_USER.value:
        fir = await mongo_storage.get_fir(case.get("fir_id"))
        if fir and fir.get("complainant_email") != current_user["email"]:
            raise HTTPException(status_code=403, detail="Access denied")
    
    if format == "pdf":
        response = await generate_pdf_report(report)
    elif format == "docx":
        response = await generate_docx_report(report)
    else:
        response = await generate_txt_report(report)
    
    response.headers["Access-Control-Allow-Origin"] = "*"
    response.headers["Access-Control-Allow-Methods"] = "GET, OPTIONS"
    response.headers["Access-Control-Allow-Headers"] = "Authorization, Content-Type"
    
    return response

@router.options("/report/download/{report_id}")
async def download_report_options():
    """Handle CORS preflight requests"""
    return Response(
        content="",
        media_type="text/plain",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "GET, OPTIONS",
            "Access-Control-Allow-Headers": "Authorization, Content-Type",
            "Access-Control-Max-Age": "86400",
        }
    )

async def generate_pdf_report(report: dict):
    """Generate PDF report with colors and formatting"""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#6366f1'),
        alignment=TA_CENTER,
        spaceAfter=30
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=12,
        spaceBefore=20
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.black,
        spaceAfter=10
    )
    
    story = []
    
    story.append(Paragraph("Forensic Analysis Report", title_style))
    story.append(Spacer(1, 12))
    
    data = [
        ["Report Number:", report.get("report_number", "N/A")],
        ["Report ID:", report.get("report_id", "N/A")],
        ["Evidence ID:", report.get("evidence_id", "N/A")],
        ["Case ID:", report.get("case_id", "N/A")],
        ["Analysis Type:", report.get("analysis_type", "N/A")],
        ["Analyst:", report.get("analyst_name", "N/A")],
        ["Created:", report.get("created_at", "N/A")],
        ["Status:", report.get("status", "N/A")]
    ]
    
    table = Table(data, colWidths=[120, 350])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#374151')),
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Findings", heading_style))
    story.append(Paragraph(report.get("findings", "No findings available"), body_style))
    story.append(Spacer(1, 15))
    
    story.append(Paragraph("Conclusion", heading_style))
    story.append(Paragraph(report.get("conclusion", "No conclusion available"), body_style))
    
    doc.build(story)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Content-Disposition": f"attachment; filename=forensic_report_{report.get('report_number')}.pdf"
        }
    )

async def generate_docx_report(report: dict):
    """Generate DOCX report with colors and formatting"""
    buffer = io.BytesIO()
    doc = docx.Document()
    
    title = doc.add_heading('Forensic Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(99, 102, 241)
    
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Report Number:", report.get("report_number", "N/A")),
        ("Report ID:", report.get("report_id", "N/A")),
        ("Evidence ID:", report.get("evidence_id", "N/A")),
        ("Case ID:", report.get("case_id", "N/A")),
        ("Analysis Type:", report.get("analysis_type", "N/A")),
        ("Analyst:", report.get("analyst_name", "N/A")),
        ("Created:", report.get("created_at", "N/A")),
        ("Status:", report.get("status", "N/A"))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(55, 65, 81)
    
    doc.add_paragraph()
    
    heading1 = doc.add_heading('Findings', level=1)
    for run in heading1.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)
    doc.add_paragraph(report.get("findings", "No findings available"))
    
    doc.add_paragraph()
    
    heading2 = doc.add_heading('Conclusion', level=1)
    for run in heading2.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)
    doc.add_paragraph(report.get("conclusion", "No conclusion available"))
    
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Content-Disposition": f"attachment; filename=forensic_report_{report.get('report_number')}.docx"
        }
    )

async def generate_txt_report(report: dict):
    """Generate plain text report"""
    content = f"""
{'='*60}
FORENSIC ANALYSIS REPORT
{'='*60}

Report Number: {report.get("report_number", "N/A")}
Report ID: {report.get("report_id", "N/A")}
Evidence ID: {report.get("evidence_id", "N/A")}
Case ID: {report.get("case_id", "N/A")}
Analysis Type: {report.get("analysis_type", "N/A")}
Analyst: {report.get("analyst_name", "N/A")}
Created: {report.get("created_at", "N/A")}
Status: {report.get("status", "N/A")}

{'='*60}
FINDINGS
{'='*60}

{report.get("findings", "No findings available")}

{'='*60}
CONCLUSION
{'='*60}

{report.get("conclusion", "No conclusion available")}

{'='*60}
Generated by Justice System
{'='*60}
"""
    return Response(
        content=content.encode(),
        media_type="text/plain",
        headers={
            "Access-Control-Allow-Origin": "*",
            "Content-Disposition": f"attachment; filename=forensic_report_{report.get('report_number')}.txt"
        }
    )

class ShareReportRequest(BaseModel):
    report_id: str
    share_with_email: str
    share_with_role: str

@router.post("/report/share")
async def share_forensic_report(
    payload: ShareReportRequest,
    current_user: dict = Depends(require_roles(UserRole.FORENSIC_ANALYST))
):
    """Share forensic report with investigator, court, or police"""
    forensic_reports = await mongo_storage.get_forensic_reports()
    report = forensic_reports.get(payload.report_id)
    
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    
    target_user = await mongo_storage.get_user(payload.share_with_email)
    if not target_user:
        raise HTTPException(status_code=404, detail="User not found")
    
    if "shared_with" not in report:
        report["shared_with"] = []
    
    for shared in report["shared_with"]:
        if shared["email"] == payload.share_with_email:
            raise HTTPException(status_code=400, detail="Already shared with this user")
    
    share_timestamp = datetime.now(timezone.utc).isoformat()
    
    share_record = {
        "email": payload.share_with_email,
        "name": target_user.get("full_name", payload.share_with_email),
        "role": payload.share_with_role,
        "shared_at": share_timestamp,
        "shared_by": current_user["email"],
        "shared_by_name": current_user.get("full_name")
    }
    report["shared_with"].append(share_record)
    
    # ADD TIMELINE EVENT TO CASE
    case = await mongo_storage.get_case(report.get("case_id"))
    if case:
        if "timeline" not in case:
            case["timeline"] = []
        
        case["timeline"].append({
            "timestamp": share_timestamp,
            "action": "📋 Forensic Report Shared",
            "description": f"Forensic report {report.get('report_number')} shared with {target_user.get('full_name')} ({payload.share_with_role})",
            "icon": "📋",
            "status": "completed",
            "by": current_user.get("full_name", current_user["email"]),
            "by_name": current_user.get("full_name", current_user["email"]),
            "type": "evidence_level",
            "details": {
                "report_id": payload.report_id,
                "report_number": report.get("report_number"),
                "shared_with_email": payload.share_with_email,
                "shared_with_name": target_user.get("full_name"),
                "shared_with_role": payload.share_with_role,
                "evidence_id": report.get("evidence_id")
            }
        })
        
        if "shared_reports" not in case:
            case["shared_reports"] = []
        case["shared_reports"].append({
            "report_id": payload.report_id,
            "report_number": report.get("report_number"),
            "shared_with": payload.share_with_email,
            "shared_with_name": target_user.get("full_name"),
            "shared_at": share_timestamp,
            "shared_by": current_user["email"],
            "shared_by_name": current_user.get("full_name")
        })
        
        await mongo_storage.save_case(report.get("case_id"), case)
        
        await manager.send_case_update(
            report.get("case_id"), 
            "timeline_update", 
            {
                "action": "report_shared",
                "report_number": report.get("report_number"),
                "shared_with": target_user.get("full_name"),
                "shared_by": current_user.get("full_name"),
                "timestamp": share_timestamp
            }
        )
    
    forensic_reports[payload.report_id] = report
    await mongo_storage.save_forensic_reports(forensic_reports)
    
    return {
        "message": f"Report shared with {target_user.get('full_name')}",
        "share_record": share_record
    }

@router.get("/report/shared/{case_id}")
async def get_shared_reports_for_case(
    case_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Get reports shared with current user for a case"""
    forensic_reports = await mongo_storage.get_forensic_reports()
    shared_reports = []
    
    for report_id, report in forensic_reports.items():
        if report.get("case_id") != case_id:
            continue
        
        for shared in report.get("shared_with", []):
            if shared["email"] == current_user["email"]:
                shared_reports.append({
                    "report_id": report_id,
                    "report_number": report.get("report_number"),
                    "analysis_type": report.get("analysis_type"),
                    "methodology": report.get("methodology", ""), 
                    "equipment_used": report.get("equipment_used", ""), 
                    "findings": report.get("findings"),
                    "conclusion": report.get("conclusion"),
                    "shared_at": shared.get("shared_at"),
                    "shared_by": shared.get("shared_by_name"),
                    "created_at": report.get("created_at")
                })
                break
    
    return shared_reports